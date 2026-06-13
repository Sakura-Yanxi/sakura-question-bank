from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(__file__).resolve().parent

DOCS = [
    ("01_软著申请信息草稿.md", "Sakura_软著申请信息草稿.docx"),
    ("02_软件设计说明书.md", "Sakura_软件设计说明书.docx"),
    ("03_用户操作手册.md", "Sakura_用户操作手册.docx"),
]


def set_run_font(run, name: str = "宋体", size: float = 10.5, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_font(paragraph, name: str = "宋体", size: float = 10.5) -> None:
    for run in paragraph.runs:
        set_run_font(run, name, size, run.bold)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, "宋体", 9.5, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "1F4D78", 14, 8),
        ("Heading 2", 13, "2E74B5", 10, 6),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    return doc


def add_title(doc: Document, title: str, subtitle: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title)
    set_run_font(run, "微软雅黑", 18, True)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(14)
        run2 = p2.add_run(subtitle)
        set_run_font(run2, "宋体", 10.5)
        run2.font.color.rgb = RGBColor.from_string("666666")


def add_metadata_table(doc: Document, kind: str) -> None:
    rows = [
        ("软件名称", "Sakura 做题集智能错题复盘系统"),
        ("版本号", "V1.0"),
        ("文档类型", kind),
        ("著作权人", "[待填写：真实姓名或单位名称]"),
        ("开发完成日期", "[待填写：以 V1.0 定稿日期为准]"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.cell(i, 0), k, True)
        set_cell_text(table.cell(i, 1), v)
        set_cell_shading(table.cell(i, 0), "F2F4F7")
    doc.add_paragraph()


def add_markdown_inline(paragraph, text: str) -> None:
    # Lightweight inline handling for `code` and **bold**.
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token.strip("`"))
            set_run_font(run, "Consolas", 9.5)
            run.font.color.rgb = RGBColor.from_string("374151")
        else:
            run = paragraph.add_run(token.strip("*"))
            set_run_font(run, "宋体", 10.5, True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line if line else " ")
        set_run_font(run, "Consolas", 8.5)
        run.font.color.rgb = RGBColor.from_string("1F2937")


def markdown_to_docx(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = lines[0].lstrip("# ").strip() if lines else md_path.stem
    doc = setup_document()
    add_title(doc, title, "计算机软件著作权登记材料草稿")
    add_metadata_table(doc, "申请信息草稿" if "申请信息" in title else ("软件设计说明书" if "设计" in title else "用户操作手册"))

    in_code = False
    code_lines: list[str] = []
    skip_first_title = True
    for raw in lines:
        line = raw.rstrip()
        if skip_first_title and line.startswith("# "):
            skip_first_title = False
            continue
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:].strip(), style="Heading 3")
            set_paragraph_font(p, "微软雅黑", 11.5)
        elif line.startswith("## "):
            p = doc.add_paragraph(line[3:].strip(), style="Heading 2")
            set_paragraph_font(p, "微软雅黑", 13)
        elif line.startswith("# "):
            p = doc.add_paragraph(line[2:].strip(), style="Heading 1")
            set_paragraph_font(p, "微软雅黑", 16)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_markdown_inline(p, line[2:].strip())
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_markdown_inline(p, re.sub(r"^\d+\.\s+", "", line).strip())
        else:
            p = doc.add_paragraph()
            add_markdown_inline(p, line)
    if code_lines:
        add_code_block(doc, code_lines)
    doc.save(out_path)


def build_source_docx() -> None:
    lines = (BASE / "04_源代码鉴别材料_分页版.txt").read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Consolas"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    normal.font.size = Pt(7.5)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        if line.startswith("===== 第") and p._p.getprevious() is not None:
            p.paragraph_format.page_break_before = True
        run = p.add_run(line if line else " ")
        set_run_font(run, "Consolas", 7.5)
        if line.startswith("===== 第"):
            run.bold = True
    doc.save(BASE / "Sakura_源代码鉴别材料_分页版.docx")


def main() -> None:
    for src, out in DOCS:
        markdown_to_docx(BASE / src, BASE / out)
    build_source_docx()
    print("docx_build_ok")


if __name__ == "__main__":
    main()
